#!/usr/bin/env python3
"""
Combined Psychology Research Survey - Word Document Generator
"The Role of Grit as a Mediator in the Relationship between Wisdom and PERMA Well-being"

Creates a Microsoft Word document (.docx) version of the combined survey
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
from docx.shared import RGBColor
from datetime import datetime
import os

class CombinedPsychologyResearchDocxGenerator:
    def __init__(self):
        # Wisdom Questionnaire questions
        self.wisdom_questions = [
            "I enjoy creating things that are new and different.",
            "I do not have many questions.",
            "I consider the positives and negatives of every option when I am making a decision.",
            "If there is a chance to learn something new, I jump right in.",
            "Others tell me that I give good advice.",
            "I see myself as a very creative person.",
            "I am curious about how things work.",
            "I carefully think about the opinions of others before I make a decision.",
            "I get excited when I see there is something new to learn.",
            "My friends ask for my opinion before they make an important decision.",
            "I often figure out different ways of doing things.",
            "I frequently ask questions.",
            "I wait until I have all the facts before I make a decision.",
            "I love learning about how to do different things.",
            "People tell me that I am a wise person.",
            "I always like to do things in different ways.",
            "I am always full of questions.",
            "I think about all my choices before I make a decision.",
            "When I want to learn something, I try to find out everything about it.",
            "I am able to solve problems in a way that is pleasing to everyone."
        ]
        
        # Grit Scale questions
        self.grit_questions = [
            "I have overcome setbacks to conquer an important challenge.",
            "New ideas and projects sometimes distract me from previous ones.",
            "My interests change from year to year.",
            "Setbacks don't discourage me. I don't give up easily.",
            "I often set a goal but later choose to pursue a different one.",
            "I have difficulty maintaining my focus on projects that take more than a few months to complete.",
            "I finish whatever I begin.",
            "I am a hard worker.",
            "I am diligent. I never give up.",
            "I have been obsessed with a certain idea or project for a short time but later lost interest.",
            "I work hard to achieve my goals.",
            "I often find myself having difficulty sticking with long-term commitments."
        ]
        
        # PERMA-Profiler questions organized by blocks
        self.perma_blocks = {
            'Block 1': {
                'questions': [
                    ('A1', 'How much of the time do you feel you are making progress towards accomplishing your goals?'),
                    ('E1', 'How often do you become absorbed in what you are doing?'),
                    ('P1', 'In general, how often do you feel joyful?'),
                    ('N1', 'In general, how often do you feel anxious?'),
                    ('A2', 'How often do you achieve the important goals you have set for yourself?')
                ],
                'scale': '0 = never, 10 = always'
            },
            'Block 2': {
                'questions': [
                    ('H1', 'In general, how would you say your health is?')
                ],
                'scale': '0 = terrible, 10 = excellent'
            },
            'Block 3': {
                'questions': [
                    ('M1', 'In general, to what extent do you lead a purposeful and meaningful life?'),
                    ('R1', 'To what extent do you receive help and support from others when you need it?'),
                    ('M2', 'In general, to what extent do you feel that what you do in your life is valuable and worthwhile?'),
                    ('E2', 'In general, to what extent do you feel excited and interested in things?'),
                    ('Lon', 'How lonely do you feel in your daily life?')
                ],
                'scale': '0 = not at all, 10 = completely'
            },
            'Block 4': {
                'questions': [
                    ('H2', 'How satisfied are you with your current physical health?')
                ],
                'scale': '0 = not at all, 10 = completely'
            },
            'Block 5': {
                'questions': [
                    ('P2', 'In general, how often do you feel positive?'),
                    ('N2', 'In general, how often do you feel angry?'),
                    ('A3', 'How often are you able to handle your responsibilities?'),
                    ('N3', 'In general, how often do you feel sad?'),
                    ('E3', 'How often do you lose track of time while doing something you enjoy?')
                ],
                'scale': '0 = never, 10 = always'
            },
            'Block 6': {
                'questions': [
                    ('H3', 'Compared to others of your same age and sex, how is your health?')
                ],
                'scale': '0 = terrible, 10 = excellent'
            },
            'Block 7': {
                'questions': [
                    ('R2', 'To what extent do you feel loved?'),
                    ('M3', 'To what extent do you generally feel you have a sense of direction in your life?'),
                    ('R3', 'How satisfied are you with your personal relationships?'),
                    ('P3', 'In general, to what extent do you feel contented?')
                ],
                'scale': '0 = not at all, 10 = completely'
            },
            'Block 8': {
                'questions': [
                    ('hap', 'Taking all things together, how happy would you say you are?')
                ],
                'scale': '0 = not at all, 10 = completely'
            }
        }

    def add_colored_paragraph(self, doc, text, font_size=12, bold=False, color=None, alignment=WD_ALIGN_PARAGRAPH.LEFT):
        """Add a paragraph with specified formatting"""
        paragraph = doc.add_paragraph()
        paragraph.alignment = alignment
        
        run = paragraph.add_run(text)
        run.font.size = Pt(font_size)
        run.bold = bold
        
        if color:
            run.font.color.rgb = color
            
        return paragraph

    def add_title_section(self, doc):
        """Add the main title and subtitle"""
        # Main title - three lines to match PDF
        title1 = doc.add_paragraph()
        title1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run1 = title1.add_run("THE ROLE OF GRIT AS A MEDIATOR IN THE")
        run1.font.size = Pt(15)
        run1.bold = True
        run1.font.color.rgb = RGBColor(44, 62, 80)  # Dark blue-grey
        
        title2 = doc.add_paragraph()
        title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = title2.add_run("RELATIONSHIP BETWEEN WISDOM AND PERMA")
        run2.font.size = Pt(15)
        run2.bold = True
        run2.font.color.rgb = RGBColor(44, 62, 80)
        
        title3 = doc.add_paragraph()
        title3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run3 = title3.add_run("WELL-BEING")
        run3.font.size = Pt(15)
        run3.bold = True
        run3.font.color.rgb = RGBColor(44, 62, 80)
        
        # Subtitle
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_run = subtitle.add_run("Research Survey")
        sub_run.font.size = Pt(11)
        sub_run.italic = True
        sub_run.font.color.rgb = RGBColor(52, 152, 219)  # Bright blue

    def add_demographic_section(self, doc):
        """Add demographic information section"""
        doc.add_paragraph()  # Space
        
        # Demographic title
        demo_title = doc.add_paragraph()
        demo_run = demo_title.add_run("DEMOGRAPHIC INFORMATION")
        demo_run.font.size = Pt(10)
        demo_run.bold = True
        
        doc.add_paragraph()  # Space
        
        # Add demographic fields as simple paragraphs instead of table
        # Name and Date
        name_para = doc.add_paragraph()
        name_run = name_para.add_run("Participant Name: _________________________     Date: __________________")
        name_run.font.size = Pt(11)
        name_para.space_after = Pt(12)
        
        # Age and Gender
        age_para = doc.add_paragraph()
        age_run = age_para.add_run("Age: _______     Gender: _________________")
        age_run.font.size = Pt(11)
        age_para.space_after = Pt(12)
        
        # Socio-Economic Status - all on one line
        ses_para = doc.add_paragraph()
        ses_run = ses_para.add_run("Socio-Economic Status:     ☐ Upper     ☐ Middle     ☐ Lower")
        ses_run.font.size = Pt(11)
        ses_para.space_after = Pt(12)
        ses_para.keep_together = True
        
        # Area - all on one line
        area_para = doc.add_paragraph()
        area_run = area_para.add_run("Area:     ☐ Rural     ☐ Urban")
        area_run.font.size = Pt(11)
        area_para.space_after = Pt(12)
        area_para.keep_together = True
        
        # Education Level - all on one line
        edu_para = doc.add_paragraph()
        edu_run = edu_para.add_run("Education Level:     ☐ High School     ☐ Intermediate     ☐ Under Graduate     ☐ Other")
        edu_run.font.size = Pt(11)
        edu_para.space_after = Pt(12)
        edu_para.keep_together = True

    def add_instructions_section(self, doc):
        """Add instructions section with proper formatting"""
        doc.add_paragraph()  # Space
        
        # Create a professional instructions box to match PDF
        inst_table = doc.add_table(rows=1, cols=1)
        inst_cell = inst_table.cell(0, 0)
        
        # Clear the default paragraph
        inst_cell.paragraphs[0].clear()
        
        # Add main instructions paragraph
        main_para = inst_cell.paragraphs[0]
        main_run = main_para.add_run("GENERAL INSTRUCTIONS: ")
        main_run.bold = True
        main_run.font.size = Pt(10)
        
        main_text_run = main_para.add_run("This survey consists of three parts that measure different psychological traits. Please read each statement carefully and respond honestly. There are no right or wrong answers - we are interested in your genuine thoughts and feelings.")
        main_text_run.font.size = Pt(10)
        main_para.space_after = Pt(8)
        
        # Add survey structure paragraph
        structure_para = inst_cell.add_paragraph()
        structure_title_run = structure_para.add_run("Part 1: ")
        structure_title_run.bold = True
        structure_title_run.font.size = Pt(10)
        
        structure_text_run = structure_para.add_run("Wisdom Assessment - 20 questions using a 5-point scale")
        structure_text_run.font.size = Pt(10)
        structure_para.space_after = Pt(4)
        
        structure_para2 = inst_cell.add_paragraph()
        structure_title_run2 = structure_para2.add_run("Part 2: ")
        structure_title_run2.bold = True
        structure_title_run2.font.size = Pt(10)
        
        structure_text_run2 = structure_para2.add_run("Grit Scale - 12 questions using a 5-point scale")
        structure_text_run2.font.size = Pt(10)
        structure_para2.space_after = Pt(4)
        
        structure_para3 = inst_cell.add_paragraph()
        structure_title_run3 = structure_para3.add_run("Part 3: ")
        structure_title_run3.bold = True
        structure_title_run3.font.size = Pt(10)
        
        structure_text_run3 = structure_para3.add_run("PERMA Well-being Assessment - 23 questions using a 0-10 scale")
        structure_text_run3.font.size = Pt(10)
        structure_para3.space_after = Pt(12)
        
        # Add completion instructions paragraph
        completion_para = inst_cell.add_paragraph()
        completion_run = completion_para.add_run("Please complete all sections and mark only one response per question. Thank you for your participation!")
        completion_run.font.size = Pt(10)
        completion_run.bold = True
        
        # Style the instructions table
        inst_table.style = 'Table Grid'
        
        # Set table cell styling
        from docx.oxml.shared import OxmlElement
        from docx.oxml.ns import qn
        
        # Add some padding to the cell
        tc = inst_cell._tc
        tcPr = tc.get_or_add_tcPr()
        
        # Set cell margins
        tcMar = OxmlElement('w:tcMar')
        
        for margin_name in ['top', 'left', 'bottom', 'right']:
            margin = OxmlElement(f'w:{margin_name}')
            margin.set(qn('w:w'), '150')  # Reduced padding
            margin.set(qn('w:type'), 'dxa')
            tcMar.append(margin)
        
        tcPr.append(tcMar)
        
        # Set background color to match PDF light blue
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), 'EBF2F9')  # Light blue background to match PDF
        tcPr.append(shd)
        
        doc.add_paragraph()  # Space after instructions

    def add_wisdom_section(self, doc):
        """Add wisdom assessment section"""
        doc.add_page_break()
        
        # Section title
        section_title = doc.add_paragraph()
        section_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = section_title.add_run("PART 1: WISDOM ASSESSMENT")
        title_run.font.size = Pt(13)
        title_run.bold = True
        
        doc.add_paragraph()
        
        # Instructions
        inst_para = doc.add_paragraph()
        inst_run = inst_para.add_run("Instructions: ")
        inst_run.bold = True
        inst_run.font.size = Pt(11)
        inst_text = inst_para.add_run("Rate each statement on how much it describes you: 5=Very much like me, 4=Mostly like me, 3=Somewhat like me, 2=A little like me, 1=Not like me at all")
        inst_text.font.size = Pt(11)
        
        doc.add_paragraph()
        
        # Add questions in table format for better layout control
        for i, question in enumerate(self.wisdom_questions, 1):
            # Create a table for each question (2 columns: question | options)
            question_table = doc.add_table(rows=1, cols=2)
            question_table.style = 'Table Grid'
            
            # Set column widths (60% for question, 40% for options)
            question_table.columns[0].width = Inches(4.0)  # Question column
            question_table.columns[1].width = Inches(2.5)  # Options column
            
            # Left cell - Question
            left_cell = question_table.cell(0, 0)
            left_para = left_cell.paragraphs[0]
            left_para.clear()
            
            q_run = left_para.add_run(f"W{i}. ")
            q_run.bold = True
            q_run.font.size = Pt(11)
            q_text = left_para.add_run(f"{question}")
            q_text.font.size = Pt(11)
            
            # Right cell - Options  
            right_cell = question_table.cell(0, 1)
            right_para = right_cell.paragraphs[0]
            right_para.clear()
            right_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            response_run = right_para.add_run("☐ 5   ☐ 4   ☐ 3   ☐ 2   ☐ 1")
            response_run.font.size = Pt(12)  # Increased font size for better visibility
            response_run.bold = True
            
            # Style the table
            question_table.style = 'Light Grid'
            
            # Add some spacing after each question
            doc.add_paragraph().space_after = Pt(6)

    def add_grit_section(self, doc):
        """Add grit scale section"""
        doc.add_page_break()
        
        # Section title
        section_title = doc.add_paragraph()
        section_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = section_title.add_run("PART 2: GRIT SCALE")
        title_run.font.size = Pt(13)
        title_run.bold = True
        
        doc.add_paragraph()
        
        # Instructions
        inst_para = doc.add_paragraph()
        inst_run = inst_para.add_run("Instructions: ")
        inst_run.bold = True
        inst_run.font.size = Pt(11)
        inst_text = inst_para.add_run("Rate each statement: 5=Very much like me, 4=Mostly like me, 3=Somewhat like me, 2=Not much like me, 1=Not like me at all")
        inst_text.font.size = Pt(11)
        
        doc.add_paragraph()
        
        # Add questions in table format for better layout control
        for i, question in enumerate(self.grit_questions, 1):
            # Create a table for each question (2 columns: question | options)
            question_table = doc.add_table(rows=1, cols=2)
            question_table.style = 'Table Grid'
            
            # Set column widths (60% for question, 40% for options)
            question_table.columns[0].width = Inches(4.0)  # Question column
            question_table.columns[1].width = Inches(2.5)  # Options column
            
            # Left cell - Question
            left_cell = question_table.cell(0, 0)
            left_para = left_cell.paragraphs[0]
            left_para.clear()
            
            q_run = left_para.add_run(f"G{i}. ")
            q_run.bold = True
            q_run.font.size = Pt(11)
            q_text = left_para.add_run(f"{question}")
            q_text.font.size = Pt(11)
            
            # Right cell - Options  
            right_cell = question_table.cell(0, 1)
            right_para = right_cell.paragraphs[0]
            right_para.clear()
            right_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            response_run = right_para.add_run("☐ 5   ☐ 4   ☐ 3   ☐ 2   ☐ 1")
            response_run.font.size = Pt(12)  # Increased font size for better visibility
            response_run.bold = True
            
            # Style the table
            question_table.style = 'Light Grid'
            
            # Add some spacing after each question
            doc.add_paragraph().space_after = Pt(6)

    def add_perma_section(self, doc):
        """Add PERMA well-being section"""
        doc.add_page_break()
        
        # Section title
        section_title = doc.add_paragraph()
        section_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = section_title.add_run("PART 3: PERMA WELL-BEING ASSESSMENT")
        title_run.font.size = Pt(13)
        title_run.bold = True
        
        doc.add_paragraph()
        
        # Instructions
        inst_para = doc.add_paragraph()
        inst_run = inst_para.add_run("Instructions: ")
        inst_run.bold = True
        inst_run.font.size = Pt(11)
        inst_text = inst_para.add_run("Rate each question on a scale from 0 to 10, where the meaning of 0 and 10 varies by question block as indicated below.")
        inst_text.font.size = Pt(11)
        
        doc.add_paragraph()
        
        # Scale header
        scale_para = doc.add_paragraph()
        scale_run = scale_para.add_run("Response: ")
        scale_run.bold = True
        scale_run.font.size = Pt(10)
        scale_text = scale_para.add_run("0   1   2   3   4   5   6   7   8   9   10")
        scale_text.font.size = Pt(8)
        scale_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # Add all PERMA question blocks
        for block_name, block_data in self.perma_blocks.items():
            # Block scale reference
            scale_ref_para = doc.add_paragraph()
            scale_ref_run = scale_ref_para.add_run(f"{block_data['scale']}")
            scale_ref_run.italic = True
            scale_ref_run.font.size = Pt(11)
            scale_ref_run.font.color.rgb = RGBColor(39, 174, 96)  # Green color to match PDF
            
            # Add questions in this block using table format
            for label, question in block_data['questions']:
                # Create a table for each question (2 columns: question | options)
                question_table = doc.add_table(rows=1, cols=2)
                question_table.style = 'Table Grid'
                
                # Set column widths for PERMA (50% for question, 50% for options - more space for 0-10 scale)
                question_table.columns[0].width = Inches(3.5)  # Question column
                question_table.columns[1].width = Inches(3.5)  # Options column (wider for 0-10 scale)
                
                # Left cell - Question
                left_cell = question_table.cell(0, 0)
                left_para = left_cell.paragraphs[0]
                left_para.clear()
                
                q_run = left_para.add_run(f"P{label}. ")
                q_run.bold = True
                q_run.font.size = Pt(11)
                q_text = left_para.add_run(f"{question}")
                q_text.font.size = Pt(11)
                
                # Right cell - Options (0-10 scale) with better spacing
                right_cell = question_table.cell(0, 1)
                right_para = right_cell.paragraphs[0]
                right_para.clear()
                right_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                response_options = "  ".join([f"☐{i}" for i in range(11)])
                response_run = right_para.add_run(response_options)
                response_run.font.size = Pt(11)  # Increased font size for better visibility
                response_run.bold = True
                
                # Style the table
                question_table.style = 'Light Grid'
                
                # Add some spacing after each question
                doc.add_paragraph().space_after = Pt(4)
            
            doc.add_paragraph()  # Space after each block

    def add_scoring_section(self, doc):
        """Add scoring section for researchers"""
        doc.add_page_break()
        
        # Section title
        section_title = doc.add_paragraph()
        section_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = section_title.add_run("FOR RESEARCH USE ONLY - SCORING SUMMARY")
        title_run.font.size = Pt(13)
        title_run.bold = True
        
        doc.add_paragraph()
        
        # Wisdom scores
        wisdom_para = doc.add_paragraph()
        wisdom_run = wisdom_para.add_run("WISDOM SCORES:")
        wisdom_run.bold = True
        wisdom_run.font.size = Pt(11)
        
        score_para = doc.add_paragraph()
        score_run = score_para.add_run("Total: ____/100 | Creativity: ____/20 | Curiosity: ____/40 | Judgment: ____/20 | Social: ____/20")
        score_run.font.size = Pt(10)
        doc.add_paragraph()
        
        # Grit scores
        grit_para = doc.add_paragraph()
        grit_run = grit_para.add_run("GRIT SCORES:")
        grit_run.bold = True
        grit_run.font.size = Pt(11)
        
        grit_score_para = doc.add_paragraph()
        grit_score_run = grit_score_para.add_run("Total: ____/60 | Consistency of Interest: ____/30 | Perseverance of Effort: ____/30")
        grit_score_run.font.size = Pt(10)
        
        reverse_para = doc.add_paragraph()
        reverse_run = reverse_para.add_run("Reverse score items G2, G3, G5, G6, G10, G12 for Grit calculations")
        reverse_run.font.size = Pt(10)
        doc.add_paragraph()
        
        # PERMA scores
        perma_para = doc.add_paragraph()
        perma_run = perma_para.add_run("PERMA SCORES:")
        perma_run.bold = True
        perma_run.font.size = Pt(11)
        
        perma_score_para = doc.add_paragraph()
        perma_score_run = perma_score_para.add_run("P (Positive): ____ | E (Engagement): ____ | R (Relationships): ____ | M (Meaning): ____ | A (Achievement): ____")
        perma_score_run.font.size = Pt(10)
        
        health_para = doc.add_paragraph()
        health_run = health_para.add_run("Health: ____ | Negative Emotions: ____ | Loneliness: ____ | Overall Happiness: ____")
        health_run.font.size = Pt(10)
        doc.add_paragraph()
        
        # Thank you
        thank_para = doc.add_paragraph()
        thank_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        thank_run = thank_para.add_run("Thank you for participating in this comprehensive well-being research study.")
        thank_run.italic = True
        thank_run.font.size = Pt(10)

    def generate_docx(self, filename="combined_psychology_research_survey.docx"):
        """Generate the complete Word document"""
        doc = Document()
        
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.8)
            section.bottom_margin = Inches(0.8)
            section.left_margin = Inches(0.8)
            section.right_margin = Inches(0.8)
        
        # Add all sections
        self.add_title_section(doc)
        self.add_demographic_section(doc)
        self.add_instructions_section(doc)
        self.add_wisdom_section(doc)
        self.add_grit_section(doc)
        self.add_perma_section(doc)
        self.add_scoring_section(doc)
        
        # Save document
        doc.save(filename)
        print(f"Word document generated successfully: {filename}")
        return filename

def main():
    """Main function to generate the Word document"""
    print("Generating Combined Psychology Research Survey Word Document...")
    
    generator = CombinedPsychologyResearchDocxGenerator()
    
    # Generate with timestamp in filename
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"combined_psychology_survey_{timestamp}.docx"
    
    try:
        docx_file = generator.generate_docx(filename)
        print(f"\n✅ Success! Word document created: {docx_file}")
        print("\nThe Word document includes:")
        print("- Editable title and demographic sections")
        print("- Part 1: Wisdom Assessment (20 questions, checkboxes)")
        print("- Part 2: Grit Scale (12 questions, checkboxes)")
        print("- Part 3: PERMA Well-being Assessment (23 questions, 0-10 scale)")
        print("- Colorful section headers and instructions")
        print("- Professional formatting with proper page breaks")
        print("- Comprehensive scoring section")
        print(f"- Total questions: 55 items plus demographics")
        
        # Also create a generic version
        generic_filename = "combined_psychology_research_survey.docx"
        generator.generate_docx(generic_filename)
        print(f"- Generic version also created: {generic_filename}")
        
        print("\n📋 Word Document Benefits:")
        print("- Participants can type directly into the document")
        print("- Easy to distribute via email")
        print("- Compatible with all versions of Microsoft Word")
        print("- Can be converted to Google Docs if needed")
        print("- Easier to collect and analyze digital responses")
        
    except Exception as e:
        print(f"❌ Error generating Word document: {e}")
        return False
    
    return True

if __name__ == "__main__":
    main()
